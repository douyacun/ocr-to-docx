import json
import logging

import easyocr
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
import math
import click

DIFF = 10
BASE_FONT_SIZE = 12
FONT_DIFF = 2


@click.group()
def handle():
    pass


@click.command("")
@click.argument("filepath")
@click.argument("outfile")
def docx(filepath, outfile):
    logging.basicConfig(level="ERROR")
    reader = easyocr.Reader(lang_list=['ch_sim', 'en'], gpu=False)
    result = reader.readtext(filepath)

    doc = Document()

    doc.styles['Normal'].font.name = u'宋体'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    doc.styles['Normal'].font.size = Pt(BASE_FONT_SIZE)

    lines = merge_line(result)
    click.echo(json.dumps(lines))

    for line in lines:
        p = doc.add_paragraph()
        if line["firstLineIndent"]:
            p.paragraph_format.first_line_indent = Pt(BASE_FONT_SIZE * 2)
        if line["isCenter"]:
            p.alignment = 1
        p.add_run(line["text"]).font.size = Pt(line["fontSize"])

    doc.save(outfile)


@click.command()
@click.argument("filepath")
# 1. 统一字体，判断字体大小？
# 2. 行间距
# 3. 横间距缩进
# 4. 标点符号，转中文
def debug(filepath):
    logging.basicConfig(level="DEBUG")
    reader = easyocr.Reader(lang_list=['ch_sim', 'en'], gpu=False)
    result = reader.readtext(filepath)

    # border = parse_border(result)
    # logging.info(f"border: {border}")

    # for i in range(len(result)):
    #     (bbox, text, prob) = result[i]
    #     logging.info(f'text: {text} bbox: {bbox}')

    logging.debug(json.dumps(merge_line(result)))


def merge_line(res):
    """
    合并同一行
    1. 判断是否位同一行, 同一行：topY <= preBottomY
    2. 判断是否需要换行，本行最后一个字没有达到最右侧边界
    3. 首行缩进？居中？
    """
    lineList = list()
    border = parse_border(res)
    logging.debug(f"border: {border}")
    fontList = parse_fontsize(res)
    logging.debug(f"font: {fontList}")

    line = {"text": "", "firstLineIndent": False, "fontSize": BASE_FONT_SIZE, "lineCount": 0, "isCenter": False}
    for i in range(len(res)):
        (bbox, text, prob) = res[i]
        pos = parse_pos(bbox)
        logging.debug(f"text {text} pos: {pos}")
        prePos = parse_pos(res[i - 1][0])
        if i == 0:
            line["text"] = text
            line["fontSize"] = get_fontsize(fontList, pos["height"])
            line["lineCount"] += 1
            if pos["leftX"] > border["left"]:
                line["firstLineIndent"] = True
        else:
            # 在图片中可以确认是统一行
            if pos["topY"] <= prePos["topY"] + 2 or pos["bottomY"] <= prePos["bottomY"] + 2:
                line["text"] += text
            else:
                # 确定是否需要换行
                if prePos["rightX"] >= border["right"] and pos["leftX"] <= border["left"]:
                    line["lineCount"] += 1
                    line["text"] += text
                else:
                    # 前一行入库
                    if line["lineCount"] == 1:
                        line["fontSize"] = get_fontsize(fontList, prePos["height"])
                        if math.fabs(math.fabs(prePos["leftX"] - border["left"]) - math.fabs(
                                prePos["rightX"] - border["right"])) <= 2:
                            line["isCenter"] = True
                            line["firstLineIndent"] = False
                    lineList.append(line)
                    # 开头
                    line = {"text": text, "firstLineIndent": False, "fontSize": BASE_FONT_SIZE, "lineCount": 1,
                            "isCenter": False}
                    if pos["leftX"] > border["left"]:
                        line["firstLineIndent"] = True

    if line["text"] != "":
        lineList.append(line)

    # logging.debug(lineList)
    return lineList


def get_fontsize(fontList: list, height: int):
    for v in fontList:
        if v["minHeight"] <= height <= v["maxHeight"]:
            return v["fontSize"]
    return BASE_FONT_SIZE


# 预估字体大小
def parse_fontsize(res) -> list:
    fontList = []
    heightDict = {}
    for i in range(len(res)):
        (bbox, text, prob) = res[i]
        pos = parse_pos(bbox)
        if pos["height"] in heightDict:
            heightDict[pos["height"]] += 1
        else:
            heightDict[pos["height"]] = 1

    c = 0
    baseHeight = 0
    for k in sorted(heightDict.keys()):
        if heightDict[k] > c:
            c = heightDict[k]
            baseHeight = k

    font = {"minHeight": 0, "maxHeight": baseHeight, "fontSize": BASE_FONT_SIZE}
    fontList.append(font)
    for k in sorted(heightDict.keys()):
        # 初始值
        if k > font["maxHeight"]:
            font = {"minHeight": k, "maxHeight": k + FONT_DIFF,
                    "fontSize": math.floor(k / baseHeight * BASE_FONT_SIZE)}
            fontList.append(font)
    return fontList


def parse_pos(bbox) -> dict:
    (top_left, top_right, bottom_right, bottom_left) = bbox
    res = {
        "topY": min(bottom_left[1], top_left[1]),
        "bottomY": min(bottom_left[1], bottom_right[1]),
        "leftX": min(top_left[0], bottom_left[0]),
        "rightX": min(top_right[0], bottom_right[0]),
    }
    res["width"] = res["rightX"] - res["leftX"]
    res["height"] = res["bottomY"] - res["topY"]
    return res


# 左右边界
def parse_border(res):
    border = dict()
    if len(res) == 0:
        return border

    right = 0
    left = 0
    for i in range(len(res)):
        (bbox, text, prob) = res[i]
        (top_left, top_right, bottom_right, bottom_left) = bbox
        if i == 0:
            left = top_left[0]
        right = max(right, top_right[0], bottom_right[0])
        left = min(left, top_left[0], bottom_left[0])

    border["left"] = left + DIFF
    border["right"] = right - DIFF
    return border


handle.add_command(docx)
handle.add_command(debug)

if __name__ == '__main__':
    handle()
