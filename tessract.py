import json
import math

import pytesseract
from PIL import Image
from xml.etree import ElementTree
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn

DIFF = 10
FONT_SIZE = 12


def to_pdf(filepath):
    pdf = pytesseract.image_to_pdf_or_hocr(Image.open(filepath), extension="pdf")
    with open('test.pdf', 'w+b') as f:
        f.write(pdf)


def image_horce(filepath, outfile):
    hocr = pytesseract.image_to_alto_xml(filepath, lang="chi_sim")
    paragraph = parse_hocr_xml(hocr=hocr, filepath="")
    print(json.dumps(paragraph))
    to_docx(paragraph, outfile)


def parse_hocr_xml(hocr, filepath):
    page = list()
    if hocr != "":
        root = ElementTree.fromstring(hocr)
        # root = tree.getroot()
        page = detect_element(root)

    if filepath != "":
        tree = ElementTree.parse(filepath)
        root = tree.getroot()
        page = detect_element(root)

    paragraph = merge_paragraph_line(page)
    return paragraph


#
# def to_docx(filepath):
#     # Get HOCR output
#     hocr = pytesseract.image_to_alto_xml(filepath, lang="chi_sim")
#     with open('test.xml', 'w+b') as f:
#         f.write(hocr)
#     # parse_hocr_xml(hocr=hocr, filepath="")


def to_docx(lines, outfile):
    # logging.basicConfig(level="ERROR")
    doc = Document()
    doc.styles['Normal'].font.name = u'宋体'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    doc.styles['Normal'].font.size = Pt(FONT_SIZE)

    # lines = merge_line(result)
    # click.echo(json.dumps(lines))

    for line in lines:
        p = doc.add_paragraph()
        if line["firstLineIndent"]:
            p.paragraph_format.first_line_indent = Pt(FONT_SIZE * 2)
        if line["isCenter"]:
            p.alignment = 1
        p.add_run(line["text"]).font.size = Pt(line["fontSize"])

    doc.save(outfile)


def merge_paragraph_line(page):
    # print(json.dumps(page))
    # print()
    border = parse_border(page)
    paragraphList = list()

    for p in page:
        isCenter = is_center(p, border)
        paragraph = {
            "firstLineIndent": not isCenter and first_line_indent(p, border),
            "isCenter": isCenter,
            "text": "",
            "fontSize": font_size(p, border)
        }
        for l in p["lineList"]:
            paragraph["text"] += l["text"]
        paragraphList.append(paragraph)
    return paragraphList


def is_center(paragraph, border):
    if len(paragraph["lineList"]) != 1:
        return False

    line = paragraph["lineList"][0]

    leftDiff = math.fabs(border["left"] - line["left"])
    rightDiff = math.fabs(border["right"] - line["right"])

    if leftDiff <= DIFF or rightDiff <= DIFF:
        return False

    if math.fabs(leftDiff - rightDiff) <= DIFF:
        return True


def first_line_indent(paragraph, border):
    if len(paragraph["lineList"]) < 1:
        return False
    firstLine = paragraph["lineList"][0]
    if math.fabs(firstLine["left"] - border["left"]) > DIFF:
        return True
    else:
        return False


def font_size(paragraph, border):
    if len(paragraph["lineList"]) != 1:
        return FONT_SIZE

    midHeight = border["midHeight"]
    line = paragraph["lineList"][0]

    if math.fabs(line["height"] - midHeight) <= DIFF:
        return FONT_SIZE

    return math.floor(line["height"] / midHeight * FONT_SIZE)


def detect_element(element: ElementTree.Element):
    ns = "http://www.loc.gov/standards/alto/ns-v3#"
    page = list()

    for composedBlock in element.find(with_prefix(ns, "Layout")). \
            find(with_prefix(ns, "Page")). \
            find(with_prefix(ns, "PrintSpace")). \
            findall(with_prefix(ns, "ComposedBlock")):
        # print(composedBlock.tag, composedBlock.attrib)
        for textBlock in composedBlock.findall(with_prefix(ns, "TextBlock")):
            # print(textBlock.tag, textBlock.attrib)
            paragraph = {"lineList": []}
            for textLine in textBlock.findall(with_prefix(ns, "TextLine")):
                # print(textLine.tag, textLine.attrib)
                line = {"text": "",
                        "height": int(textLine.attrib["HEIGHT"]),
                        "left": int(textLine.attrib["HPOS"]),
                        "right": int(textLine.attrib["HPOS"]) + int(textLine.attrib["WIDTH"])}
                for str in textLine.findall(with_prefix(ns, "String")):
                    # print(str.tag, str.attrib)
                    if str.attrib["CONTENT"] != "":
                        line["text"] += str.attrib["CONTENT"]
                paragraph["lineList"].append(line)
            page.append(paragraph)

    return page


def parse_border(page):
    minLeft = 0
    maxRight = 0
    midHeight = 0
    heightDict = dict()

    for paragraph in page:
        for line in paragraph["lineList"]:
            if minLeft == 0:
                minLeft = line["left"]

            minLeft = min(minLeft, line["left"])
            maxRight = max(maxRight, line["right"])
            if line["height"] in heightDict:
                heightDict[line["height"]] += 1
            else:
                heightDict[line["height"]] = 1
    c = 0
    for k in heightDict.keys():
        if heightDict[k] > c:
            c = heightDict[k]
            midHeight = k

    border = {
        "left": minLeft,
        "right": maxRight,
        "midHeight": midHeight
    }
    return border


def with_prefix(prefix, tag):
    return "{%s}%s" % (prefix, tag)


if __name__ == '__main__':
    image_horce("/Users/admin/Desktop/image-ocr/c.png", "temp/c.docx")
    # parse_hocr_xml(filepath="temp/test.xml", hocr="")
